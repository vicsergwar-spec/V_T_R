"""
Servicio de gestión de archivos y carpetas
"""
import os
import re
import json
import shutil
import logging
import unicodedata
from pathlib import Path
from datetime import datetime
from typing import Optional

logger = logging.getLogger(__name__)

# Detecta el sufijo de fecha/hora que añadimos al crear clases: _YYYY-MM-DD_HH-MM
_DATE_SUFFIX_RE = re.compile(r'^(.+?)_(\d{4}-\d{2}-\d{2})_(\d{2}-\d{2})$')


VIDEOS_ORIGINALES_DIR = Path(__file__).parent.parent / "videos_originales"


class FileManager:
    """Gestiona la organización de archivos y carpetas del sistema"""

    def __init__(self, base_dir: str, clases_dir: str, temp_dir: str):
        self.base_dir = Path(base_dir)
        self.clases_dir = Path(clases_dir)
        self.temp_dir = Path(temp_dir)

        self.clases_dir.mkdir(parents=True, exist_ok=True)
        self.temp_dir.mkdir(parents=True, exist_ok=True)
        VIDEOS_ORIGINALES_DIR.mkdir(parents=True, exist_ok=True)

        logger.info(f"FileManager inicializado. Clases en: {self.clases_dir}")

    # ──────────────────────────────────────────────────────────
    # Gestión de clases
    # ──────────────────────────────────────────────────────────

    def create_class_folder(self, folder_name: str, parent_path: str = "") -> Path:
        """
        Crea una carpeta para una nueva clase.
        El nombre final será: {nombre_IA}_{YYYY-MM-DD}_{HH-MM}

        Args:
            folder_name: Nombre generado por Gemini (Materia_Tema)
            parent_path: Ruta relativa de la carpeta destino (vacío = raíz)

        Returns:
            Path de la carpeta creada
        """
        folder_name = self._sanitize_folder_name(folder_name)

        # Añadir sufijo de fecha y hora
        now = datetime.now()
        date_suffix = now.strftime("%Y-%m-%d_%H-%M")
        full_name = f"{folder_name}_{date_suffix}"

        # Directorio padre
        if parent_path:
            parent_dir = self.clases_dir / parent_path
            parent_dir.mkdir(parents=True, exist_ok=True)
        else:
            parent_dir = self.clases_dir

        class_folder = parent_dir / full_name

        # Manejar duplicados
        if class_folder.exists():
            counter = 1
            while (parent_dir / f"{full_name}_{counter}").exists():
                counter += 1
            class_folder = parent_dir / f"{full_name}_{counter}"

        class_folder.mkdir(parents=True, exist_ok=True)
        logger.info(f"Carpeta de clase creada: {class_folder}")
        return class_folder

    def save_transcription(self, segments: list, class_folder: Path) -> str:
        output_path = class_folder / "transcripcion.jsonl"
        with open(output_path, "w", encoding="utf-8") as f:
            for segment in segments:
                f.write(json.dumps(segment, ensure_ascii=False) + "\n")
        logger.info(f"Transcripción guardada: {output_path}")
        return str(output_path)

    def save_summary(self, summary: str, class_folder: Path) -> str:
        output_path = class_folder / "resumen.md"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(summary)
        logger.info(f"Resumen guardado: {output_path}")
        return str(output_path)

    def save_slides(self, slides_markdown: str, class_folder: Path) -> str:
        """Guarda el contenido de slides extraídos como Markdown."""
        output_path = class_folder / "slides.md"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(slides_markdown)
        logger.info(f"Slides guardados: {output_path}")
        return str(output_path)

    def save_slides_document(self, content: str, class_folder: Path) -> str:
        """Guarda el documento de slides generado por IA (optimizado para lectura y chat)."""
        output_path = class_folder / "slides_document.md"
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"Documento de slides guardado: {output_path}")
        return str(output_path)

    def get_slides(self, class_id: str) -> Optional[str]:
        """Devuelve el contenido de slides si existe, o None."""
        slides_path = self.clases_dir / class_id / "slides.md"
        if slides_path.exists():
            return slides_path.read_text(encoding="utf-8")
        return None

    def get_slides_document(self, class_id: str) -> Optional[str]:
        """Devuelve el documento de slides generado por IA, o None."""
        doc_path = self.clases_dir / class_id / "slides_document.md"
        if doc_path.exists():
            return doc_path.read_text(encoding="utf-8")
        return None

    def save_video_to_temp(self, file_storage, filename: str) -> str:
        safe_filename = self._sanitize_filename(filename)
        temp_path = self.temp_dir / safe_filename
        file_storage.save(str(temp_path))
        logger.info(f"Video guardado temporalmente: {temp_path}")
        return str(temp_path)

    def cleanup_temp_files(self, video_path: str, audio_path: str = None, delete_video: bool = False) -> None:
        """
        Limpia archivos temporales.
        Por defecto NO elimina el video (se preserva con preserve_video).
        Pasar delete_video=True en paths de error donde no se preservó el video.
        """
        try:
            if delete_video and video_path and Path(video_path).exists():
                os.remove(video_path)
                logger.info(f"Video temporal eliminado: {video_path}")
            if audio_path and Path(audio_path).exists():
                os.remove(audio_path)
                logger.info(f"Audio temporal eliminado: {audio_path}")
        except Exception as e:
            logger.warning(f"Error al eliminar archivos temporales: {e}")

    def preserve_video(self, video_path: str, folder_name: str, folder_path: str = "") -> Optional[str]:
        """
        Mueve el video original a videos_originales/<materia>/<nombre>.ext
        en lugar de eliminarlo.

        Args:
            video_path: Ruta temporal del video
            folder_name: Nombre de la carpeta de clase (usado como nombre del video)
            folder_path: Carpeta padre (materia), ej: "Matematicas"

        Returns:
            Ruta final del video, o None si falla
        """
        try:
            src = Path(video_path)
            if not src.exists():
                return None

            # Determinar subdirectorio de materia
            materia = folder_path.replace("/", os.sep) if folder_path else "General"
            dest_dir = VIDEOS_ORIGINALES_DIR / materia
            dest_dir.mkdir(parents=True, exist_ok=True)

            # Nombre del archivo: nombre de clase + extensión original
            ext = src.suffix
            safe_name = "".join(c for c in folder_name if c.isalnum() or c in "_- ")[:80]
            dest = dest_dir / f"{safe_name}{ext}"

            # Evitar colisiones
            if dest.exists():
                counter = 1
                while (dest_dir / f"{safe_name}_{counter}{ext}").exists():
                    counter += 1
                dest = dest_dir / f"{safe_name}_{counter}{ext}"

            shutil.move(str(src), str(dest))
            logger.info(f"[I] Video original preservado: {dest}")
            return str(dest)
        except Exception as e:
            logger.warning(f"Error preservando video original: {e}")
            # Fallback: no eliminar el video, dejarlo en temp
            return None

    def get_preserved_video(self, class_id: str) -> Optional[str]:
        """
        Busca el video original preservado para una clase dada.
        Busca en videos_originales/<materia>/ por nombre de clase.
        """
        # Extraer nombre base de la clase (sin fecha)
        folder_name = class_id.split("/")[-1] if "/" in class_id else class_id
        materia = "/".join(class_id.split("/")[:-1]) if "/" in class_id else "General"

        search_dir = VIDEOS_ORIGINALES_DIR / materia.replace("/", os.sep)
        if not search_dir.exists():
            # Buscar en toda la carpeta de videos_originales
            search_dir = VIDEOS_ORIGINALES_DIR

        # Buscar archivos cuyo nombre contenga el nombre de la carpeta de clase
        for f in search_dir.rglob("*"):
            if f.is_file() and f.suffix.lower() in (
                '.mp4', '.mkv', '.avi', '.mov', '.webm',
                '.flv', '.wmv', '.m4v', '.mpeg', '.mpg'
            ):
                if folder_name in f.stem or f.stem in folder_name:
                    return str(f)
        return None

    def regenerate_slides(self, class_id: str, slide_extractor, gemini_service=None) -> dict:
        """
        Reprocesa los slides de una clase desde cero usando el video original.
        NO toca transcripcion.jsonl ni resumen.md.

        Args:
            class_id: ID de la clase (ruta relativa)
            slide_extractor: Instancia de SlideExtractor
            gemini_service: Instancia de GeminiService (para regenerar documento)

        Returns:
            dict con resultado: {success, message, slides_count}
        """
        class_folder = self.clases_dir / class_id
        if not class_folder.exists():
            return {"success": False, "message": "Clase no encontrada"}

        # Buscar video original
        video_path = self.get_preserved_video(class_id)
        if not video_path or not Path(video_path).exists():
            return {"success": False, "message": "Video original no encontrado en videos_originales/"}

        # Borrar slides e imágenes anteriores
        slides_md_path = class_folder / "slides.md"
        slides_doc_path = class_folder / "slides_document.md"
        slide_images_dir = class_folder / "slide_images"

        if slides_md_path.exists():
            slides_md_path.unlink()
            logger.info(f"[I] slides.md anterior eliminado: {class_id}")
        if slides_doc_path.exists():
            slides_doc_path.unlink()
            logger.info(f"[I] slides_document.md anterior eliminado: {class_id}")
        if slide_images_dir.exists():
            shutil.rmtree(slide_images_dir)
            logger.info(f"[I] slide_images/ anterior eliminado: {class_id}")

        # Reprocesar slides
        try:
            slides = slide_extractor.extract_slides(
                video_path=video_path,
                temp_dir=str(self.temp_dir),
                persist_dir=str(class_folder),
            )
            slides_storage_md = slide_extractor.format_slides_for_storage(slides)
            n_useful = len([s for s in slides if s.get("text") or s.get("visual_description")])

            if slides_storage_md:
                self.save_slides(slides_storage_md, class_folder)
                logger.info(f"[I] Slides regenerados: {n_useful} con contenido")

                # Regenerar documento de slides con IA si hay servicio
                if gemini_service:
                    try:
                        folder_name = class_id.split('/')[-1] if '/' in class_id else class_id
                        # Construir mapa de imágenes
                        import re as _r
                        image_map = {}
                        new_images_dir = class_folder / "slide_images"
                        if new_images_dir.exists():
                            for fp in sorted(new_images_dir.glob("slide_*.png")):
                                m = _r.match(r'slide_(\d+)(?:_sub_\d+)?\.png', fp.name)
                                if m:
                                    slide_num = int(m.group(1))
                                    image_map.setdefault(slide_num, []).append(f"slide_images/{fp.name}")
                        slides_document = gemini_service.generate_slides_document(
                            slides_storage_md, folder_name, image_map=image_map
                        )
                        if slides_document:
                            self.save_slides_document(slides_document, class_folder)
                            logger.info("[I] Documento de slides regenerado con IA")
                    except Exception as e:
                        logger.warning(f"Error regenerando documento de slides: {e}")

                return {"success": True, "message": f"{n_useful} slides regenerados", "slides_count": n_useful}
            else:
                return {"success": True, "message": "No se encontró contenido en los slides", "slides_count": 0}
        except Exception as e:
            logger.error(f"Error regenerando slides: {e}")
            return {"success": False, "message": str(e)}

    def get_all_classes(self) -> list:
        """
        Obtiene todas las clases de forma recursiva (incluyendo las de subcarpetas).
        Ordenadas de más vieja a más nueva.
        """
        classes = []
        if not self.clases_dir.exists():
            return classes
        try:
            self._scan_for_classes(self.clases_dir, "", classes)
        except OSError as e:
            logger.error(f"Error al leer directorio de clases: {e}")
        classes.sort(key=lambda x: x.get("created_at", ""))  # más vieja primero
        return classes

    def _scan_for_classes(self, directory: Path, relative_path: str, classes: list) -> None:
        """Escanea recursivamente buscando carpetas de clase (las que tienen transcripcion.jsonl)"""
        try:
            for item in sorted(directory.iterdir(), key=lambda p: p.name):
                if not item.is_dir():
                    continue
                item_rel = f"{relative_path}/{item.name}" if relative_path else item.name
                class_info = self._get_class_info(item, item_rel)
                if class_info:
                    classes.append(class_info)
                else:
                    # Es una carpeta de organización, recursear
                    self._scan_for_classes(item, item_rel, classes)
        except OSError as e:
            logger.error(f"Error escaneando {directory}: {e}")

    def get_class_by_id(self, class_id: str) -> Optional[dict]:
        """
        Obtiene información de una clase por su ID (ruta relativa desde clases_dir).

        Args:
            class_id: Ruta relativa, ej: "Matematicas/Git_2026-02-19_22-07" o "Git_2026-02-19"
        """
        class_folder = self.clases_dir / class_id
        if not class_folder.exists() or not class_folder.is_dir():
            return None
        return self._get_class_info(class_folder, class_id)

    def get_transcription(self, class_id: str) -> Optional[list]:
        transcription_path = self.clases_dir / class_id / "transcripcion.jsonl"
        if not transcription_path.exists():
            return None
        segments = []
        with open(transcription_path, "r", encoding="utf-8") as f:
            for line in f:
                if line.strip():
                    segments.append(json.loads(line))
        return segments

    def get_transcription_text(self, class_id: str) -> Optional[str]:
        segments = self.get_transcription(class_id)
        if segments:
            return " ".join(segment["texto"] for segment in segments)
        return None

    def get_summary(self, class_id: str) -> Optional[str]:
        summary_path = self.clases_dir / class_id / "resumen.md"
        if not summary_path.exists():
            return None
        with open(summary_path, "r", encoding="utf-8") as f:
            return f.read()

    def get_chat_history(self, class_id: str) -> Optional[list]:
        """Lee el historial de chat guardado en disco para una clase."""
        history_path = self.clases_dir / class_id / "chat_historial.json"
        if not history_path.exists():
            return None
        try:
            with open(history_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Error al leer historial de chat: {e}")
            return None

    def save_chat_history(self, class_id: str, history: list) -> None:
        """Guarda el historial de chat en disco."""
        history_path = self.clases_dir / class_id / "chat_historial.json"
        try:
            with open(history_path, "w", encoding="utf-8") as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"Error al guardar historial de chat: {e}")

    def delete_chat_history(self, class_id: str) -> None:
        """Elimina el historial de chat del disco."""
        history_path = self.clases_dir / class_id / "chat_historial.json"
        if history_path.exists():
            try:
                history_path.unlink()
            except Exception as e:
                logger.warning(f"Error al eliminar historial de chat: {e}")

    def get_cache_name(self, class_id: str) -> Optional[str]:
        """Lee el nombre del caché de Gemini guardado para una clase."""
        path = self.clases_dir / class_id / "gemini_cache.txt"
        if not path.exists():
            return None
        content = path.read_text(encoding="utf-8").strip()
        return content if content else None

    def save_cache_name(self, class_id: str, cache_name: str) -> None:
        """Guarda el nombre del caché de Gemini para una clase."""
        path = self.clases_dir / class_id / "gemini_cache.txt"
        path.write_text(cache_name, encoding="utf-8")

    def delete_cache_name(self, class_id: str) -> None:
        """Elimina el archivo de caché de Gemini de una clase."""
        path = self.clases_dir / class_id / "gemini_cache.txt"
        if path.exists():
            try:
                path.unlink()
            except Exception as e:
                logger.warning(f"Error al eliminar caché de clase: {e}")

    # ──────────────────────────────────────────────────────────
    # Chat general de carpeta
    # ──────────────────────────────────────────────────────────

    def get_folder_all_content(self, folder_path: str) -> list:
        """
        Recopila el contenido de TODAS las clases dentro de una carpeta (recursivo).
        Devuelve lista de dicts: {"name", "transcription", "summary", "slides"}
        """
        target_dir = self.clases_dir / folder_path
        if not target_dir.exists() or not target_dir.is_dir():
            return []

        classes = []
        self._scan_for_classes(target_dir, folder_path, classes)

        result = []
        for cls in classes:
            result.append({
                "name": cls["name"],
                "transcription": self.get_transcription_text(cls["id"]) or "",
                "summary": self.get_summary(cls["id"]) or "",
                "slides": self.get_slides_document(cls["id"]) or self.get_slides(cls["id"]) or "",
            })
        return result

    def get_folder_chat_history(self, folder_path: str) -> Optional[list]:
        """Lee el historial de chat de carpeta guardado en disco."""
        path = self.clases_dir / folder_path / "folder_chat_historial.json"
        if not path.exists():
            return None
        try:
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception as e:
            logger.warning(f"Error al leer historial de chat de carpeta: {e}")
            return None

    def save_folder_chat_history(self, folder_path: str, history: list) -> None:
        """Guarda el historial de chat de carpeta en disco."""
        path = self.clases_dir / folder_path / "folder_chat_historial.json"
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(history, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"Error al guardar historial de chat de carpeta: {e}")

    def delete_folder_chat_history(self, folder_path: str) -> None:
        """Elimina el historial de chat de carpeta del disco."""
        path = self.clases_dir / folder_path / "folder_chat_historial.json"
        if path.exists():
            try:
                path.unlink()
            except Exception as e:
                logger.warning(f"Error al eliminar historial de chat de carpeta: {e}")

    def get_folder_cache_name(self, folder_path: str) -> Optional[str]:
        """Lee el nombre del caché de Gemini para el chat de carpeta."""
        path = self.clases_dir / folder_path / "folder_gemini_cache.txt"
        if not path.exists():
            return None
        content = path.read_text(encoding="utf-8").strip()
        return content if content else None

    def save_folder_cache_name(self, folder_path: str, cache_name: str) -> None:
        """Guarda el nombre del caché de Gemini para el chat de carpeta."""
        path = self.clases_dir / folder_path / "folder_gemini_cache.txt"
        path.write_text(cache_name, encoding="utf-8")

    def delete_folder_cache_name(self, folder_path: str) -> None:
        """Elimina el archivo de caché de Gemini del chat de carpeta."""
        path = self.clases_dir / folder_path / "folder_gemini_cache.txt"
        if path.exists():
            try:
                path.unlink()
            except Exception as e:
                logger.warning(f"Error al eliminar caché de carpeta: {e}")

    def delete_class(self, class_id: str) -> bool:
        class_folder = self.clases_dir / class_id
        if not class_folder.exists():
            return False
        try:
            shutil.rmtree(class_folder)
            logger.info(f"Clase eliminada: {class_id}")
            return True
        except Exception as e:
            logger.error(f"Error al eliminar clase: {e}")
            return False

    # ──────────────────────────────────────────────────────────
    # Gestión de carpetas de organización
    # ──────────────────────────────────────────────────────────

    def get_folder_tree(self) -> list:
        """
        Devuelve lista plana de carpetas de organización (no clases)
        ordenadas por ruta.
        """
        folders = []
        if self.clases_dir.exists():
            self._scan_for_folders(self.clases_dir, "", folders)
        return sorted(folders, key=lambda x: x["path"])

    def _scan_for_folders(self, directory: Path, relative_path: str, folders: list) -> None:
        """Escanea directorios que NO son clases (no tienen transcripcion.jsonl)"""
        try:
            for item in sorted(directory.iterdir(), key=lambda p: p.name):
                if not item.is_dir():
                    continue
                # Si tiene archivos de clase, no es carpeta de organización
                if (item / "transcripcion.jsonl").exists() or (item / "resumen.md").exists():
                    continue
                item_rel = f"{relative_path}/{item.name}" if relative_path else item.name
                depth = item_rel.count("/")
                folders.append({
                    "path": item_rel,
                    "name": item.name,
                    "depth": depth
                })
                self._scan_for_folders(item, item_rel, folders)
        except OSError as e:
            logger.error(f"Error escaneando carpetas en {directory}: {e}")

    def create_folder(self, folder_path: str) -> dict:
        """
        Crea una carpeta de organización (no una clase).
        Los nombres se normalizan sin tildes para evitar problemas en rutas.

        Args:
            folder_path: Ruta relativa, ej: "Matematicas" o "Matematicas/Calculo"
        """
        parts = folder_path.replace("\\", "/").split("/")
        safe_parts = [self._sanitize_folder_name(p) for p in parts if p.strip()]
        if not safe_parts:
            raise ValueError("Nombre de carpeta inválido")
        folder = self.clases_dir / "/".join(safe_parts)
        folder.mkdir(parents=True, exist_ok=True)
        rel_path = "/".join(safe_parts)
        logger.info(f"Carpeta de organización creada: {folder}")
        return {"path": rel_path, "name": safe_parts[-1], "depth": len(safe_parts) - 1}

    def rename_class(self, class_id: str, new_name: str) -> bool:
        """
        Guarda un nombre personalizado para una clase en nombre.txt.
        No mueve ni renombra la carpeta, preservando el ID y todos los archivos.
        """
        class_folder = self.clases_dir / class_id
        if not class_folder.exists() or not class_folder.is_dir():
            return False
        new_name = new_name.strip()
        if not new_name:
            return False
        (class_folder / "nombre.txt").write_text(new_name, encoding="utf-8")
        logger.info(f"Clase renombrada: {class_id} → {new_name}")
        return True

    # ──────────────────────────────────────────────────────────
    # Conocimiento extra y rúbricas
    # ──────────────────────────────────────────────────────────

    def save_knowledge_file(self, class_id: str, filename: str, file_storage) -> str:
        """Guarda un archivo de conocimiento extra en extra_knowledge/"""
        knowledge_dir = self.clases_dir / class_id / "extra_knowledge"
        knowledge_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._sanitize_knowledge_filename(filename)
        dest = knowledge_dir / safe_name
        file_storage.save(str(dest))
        logger.info(f"Archivo de conocimiento guardado: {dest}")
        return safe_name

    def get_knowledge_files(self, class_id: str) -> list:
        """Lista archivos en extra_knowledge/"""
        knowledge_dir = self.clases_dir / class_id / "extra_knowledge"
        if not knowledge_dir.exists():
            return []
        files = []
        for f in sorted(knowledge_dir.iterdir()):
            if f.is_file():
                files.append({"name": f.name, "size": f.stat().st_size})
        return files

    def delete_knowledge_file(self, class_id: str, filename: str) -> bool:
        """Elimina un archivo de conocimiento extra."""
        fpath = self.clases_dir / class_id / "extra_knowledge" / filename
        if fpath.exists():
            fpath.unlink()
            logger.info(f"Archivo de conocimiento eliminado: {fpath}")
            return True
        return False

    def get_knowledge_text(self, class_id: str) -> str:
        """Lee todos los archivos de texto en extra_knowledge/ y devuelve texto combinado."""
        knowledge_dir = self.clases_dir / class_id / "extra_knowledge"
        if not knowledge_dir.exists():
            return ""
        parts = []
        for f in sorted(knowledge_dir.iterdir()):
            if not f.is_file():
                continue
            ext = f.suffix.lower()
            if ext in ('.txt', '.md', '.markdown'):
                try:
                    content = f.read_text(encoding="utf-8")
                    parts.append(f"--- Archivo: {f.name} ---\n{content}")
                except Exception as e:
                    logger.warning(f"Error leyendo {f}: {e}")
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    with open(f, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        text = "\n".join(page.extract_text() or "" for page in reader.pages)
                        parts.append(f"--- Archivo PDF: {f.name} ---\n{text}")
                except Exception:
                    parts.append(f"--- Archivo PDF: {f.name} (contenido no extraible) ---")
            elif ext in ('.docx', '.doc'):
                try:
                    import docx
                    doc = docx.Document(str(f))
                    text = "\n".join(p.text for p in doc.paragraphs)
                    parts.append(f"--- Archivo DOCX: {f.name} ---\n{text}")
                except Exception:
                    parts.append(f"--- Archivo DOCX: {f.name} (contenido no extraible) ---")
            elif ext in ('.png', '.jpg', '.jpeg', '.webp'):
                parts.append(f"--- Imagen: {f.name} (archivo de imagen subido) ---")
        return "\n\n".join(parts)

    def save_rubrica(self, class_id: str, filename: str, content: str) -> str:
        """Guarda una rúbrica de texto."""
        rubrica_dir = self.clases_dir / class_id / "rubricas"
        rubrica_dir.mkdir(parents=True, exist_ok=True)
        dest = rubrica_dir / filename
        dest.write_text(content, encoding="utf-8")
        logger.info(f"Rúbrica guardada: {dest}")
        return filename

    def save_rubrica_file(self, class_id: str, filename: str, file_storage) -> str:
        """Guarda un archivo de rúbrica subido."""
        rubrica_dir = self.clases_dir / class_id / "rubricas"
        rubrica_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._sanitize_knowledge_filename(filename)
        dest = rubrica_dir / safe_name
        file_storage.save(str(dest))
        logger.info(f"Archivo de rúbrica guardado: {dest}")
        return safe_name

    def get_rubrica_files(self, class_id: str) -> list:
        """Lista archivos de rúbricas."""
        rubrica_dir = self.clases_dir / class_id / "rubricas"
        if not rubrica_dir.exists():
            return []
        files = []
        for f in sorted(rubrica_dir.iterdir()):
            if f.is_file():
                files.append({"name": f.name, "size": f.stat().st_size})
        return files

    def delete_rubrica_file(self, class_id: str, filename: str) -> bool:
        """Elimina un archivo de rúbrica."""
        fpath = self.clases_dir / class_id / "rubricas" / filename
        if fpath.exists():
            fpath.unlink()
            logger.info(f"Archivo de rúbrica eliminado: {fpath}")
            return True
        return False

    def get_rubricas_text(self, class_id: str) -> str:
        """Lee todas las rúbricas y devuelve texto combinado."""
        rubrica_dir = self.clases_dir / class_id / "rubricas"
        if not rubrica_dir.exists():
            return ""
        parts = []
        for f in sorted(rubrica_dir.iterdir()):
            if not f.is_file():
                continue
            ext = f.suffix.lower()
            if ext in ('.txt', '.md', '.markdown'):
                try:
                    content = f.read_text(encoding="utf-8")
                    parts.append(f"--- Rúbrica: {f.name} ---\n{content}")
                except Exception as e:
                    logger.warning(f"Error leyendo rúbrica {f}: {e}")
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    with open(f, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        text = "\n".join(page.extract_text() or "" for page in reader.pages)
                        parts.append(f"--- Rúbrica PDF: {f.name} ---\n{text}")
                except Exception:
                    parts.append(f"--- Rúbrica PDF: {f.name} (contenido no extraible) ---")
            elif ext in ('.docx', '.doc'):
                try:
                    import docx
                    doc = docx.Document(str(f))
                    text = "\n".join(p.text for p in doc.paragraphs)
                    parts.append(f"--- Rúbrica DOCX: {f.name} ---\n{text}")
                except Exception:
                    parts.append(f"--- Rúbrica DOCX: {f.name} (contenido no extraible) ---")
        return "\n\n".join(parts)

    # ──────────────────────────────────────────────────────────
    # Imágenes de contexto para chat
    # ──────────────────────────────────────────────────────────

    def save_context_image(self, class_id: str, filename: str, file_storage) -> str:
        """Guarda una imagen de contexto en context_images/"""
        images_dir = self.clases_dir / class_id / "context_images"
        images_dir.mkdir(parents=True, exist_ok=True)
        safe_name = self._sanitize_knowledge_filename(filename)
        dest = images_dir / safe_name
        file_storage.save(str(dest))
        logger.info(f"Imagen de contexto guardada: {dest}")
        return safe_name

    def get_context_images(self, class_id: str) -> list:
        """Lista archivos de imágenes de contexto."""
        images_dir = self.clases_dir / class_id / "context_images"
        if not images_dir.exists():
            return []
        files = []
        for f in sorted(images_dir.iterdir()):
            if f.is_file() and f.suffix.lower() in ('.jpg', '.jpeg', '.png', '.webp', '.gif'):
                files.append({"name": f.name, "size": f.stat().st_size})
        return files

    def delete_context_image(self, class_id: str, filename: str) -> bool:
        """Elimina una imagen de contexto."""
        fpath = self.clases_dir / class_id / "context_images" / filename
        if fpath.exists():
            fpath.unlink()
            logger.info(f"Imagen de contexto eliminada: {fpath}")
            return True
        return False

    def get_context_images_data(self, class_id: str) -> list:
        """Lee todas las imágenes de contexto como datos binarios para enviar a Gemini."""
        images_dir = self.clases_dir / class_id / "context_images"
        if not images_dir.exists():
            return []
        mime_map = {
            '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
            '.png': 'image/png', '.webp': 'image/webp',
            '.gif': 'image/gif',
        }
        result = []
        for f in sorted(images_dir.iterdir()):
            ext = f.suffix.lower()
            if f.is_file() and ext in mime_map:
                try:
                    data = f.read_bytes()
                    result.append({
                        "name": f.name,
                        "mime_type": mime_map[ext],
                        "data": data,
                    })
                except Exception as e:
                    logger.warning(f"Error leyendo imagen de contexto {f}: {e}")
        return result

    @staticmethod
    def _sanitize_knowledge_filename(filename: str) -> str:
        """Sanitiza un nombre de archivo para conocimiento/rúbrica: sin tildes."""
        path = Path(filename)
        name = path.stem
        ext = path.suffix
        # Eliminar tildes/diacríticos
        nfkd = unicodedata.normalize("NFKD", name)
        name = "".join(c for c in nfkd if not unicodedata.combining(c))
        name = "".join(c for c in name if c.isalnum() or c in "._- ")
        name = name.replace(" ", "_")
        if not name:
            name = "archivo"
        return f"{name}{ext}"

    # ──────────────────────────────────────────────────────────
    # Extra knowledge global (carpeta extra_knowledge/ del proyecto)
    # ──────────────────────────────────────────────────────────

    def _ensure_extra_knowledge_dir(self) -> Path:
        """Devuelve la ruta de la carpeta global extra_knowledge, creándola si no existe."""
        ek_dir = self.base_dir / "extra_knowledge"
        ek_dir.mkdir(parents=True, exist_ok=True)
        return ek_dir

    def get_extra_knowledge_content(self) -> str:
        """Lee todos los archivos .txt/.md/.pdf/.docx en extra_knowledge/ y devuelve texto combinado."""
        ek_dir = self._ensure_extra_knowledge_dir()
        parts = []
        for f in sorted(ek_dir.iterdir()):
            if not f.is_file():
                continue
            ext = f.suffix.lower()
            size = f.stat().st_size
            if ext in ('.txt', '.md', '.markdown'):
                try:
                    content = f.read_text(encoding="utf-8")
                    parts.append(f"--- Archivo: {f.name} ---\n{content}")
                    logger.info(f"[I] Extra knowledge inyectado: {f.name} ({size} bytes)")
                except Exception as e:
                    logger.warning(f"Error leyendo extra knowledge {f}: {e}")
            elif ext == '.pdf':
                try:
                    import PyPDF2
                    with open(f, 'rb') as pdf_file:
                        reader = PyPDF2.PdfReader(pdf_file)
                        text = "\n".join(page.extract_text() or "" for page in reader.pages)
                        parts.append(f"--- Archivo PDF: {f.name} ---\n{text}")
                        logger.info(f"[I] Extra knowledge inyectado: {f.name} ({size} bytes)")
                except Exception:
                    parts.append(f"--- Archivo PDF: {f.name} (contenido no extraíble) ---")
            elif ext in ('.docx', '.doc'):
                try:
                    import docx
                    doc = docx.Document(str(f))
                    text = "\n".join(p.text for p in doc.paragraphs)
                    parts.append(f"--- Archivo DOCX: {f.name} ---\n{text}")
                    logger.info(f"[I] Extra knowledge inyectado: {f.name} ({size} bytes)")
                except Exception:
                    parts.append(f"--- Archivo DOCX: {f.name} (contenido no extraíble) ---")
        return "\n\n".join(parts)

    def save_extra_knowledge_file(self, filename: str, file_storage) -> str:
        """Guarda un archivo en la carpeta global extra_knowledge/."""
        ek_dir = self._ensure_extra_knowledge_dir()
        safe_name = self._sanitize_knowledge_filename(filename)
        dest = ek_dir / safe_name
        file_storage.save(str(dest))
        logger.info(f"Extra knowledge guardado: {dest} ({dest.stat().st_size} bytes)")
        return safe_name

    def save_extra_knowledge_text(self, filename: str, content: str) -> str:
        """Guarda texto como archivo en la carpeta global extra_knowledge/."""
        ek_dir = self._ensure_extra_knowledge_dir()
        dest = ek_dir / filename
        dest.write_text(content, encoding="utf-8")
        logger.info(f"Extra knowledge texto guardado: {dest} ({dest.stat().st_size} bytes)")
        return filename

    def list_extra_knowledge_files(self) -> list:
        """Lista archivos en la carpeta global extra_knowledge/."""
        ek_dir = self._ensure_extra_knowledge_dir()
        files = []
        for f in sorted(ek_dir.iterdir()):
            if f.is_file():
                files.append({"name": f.name, "size": f.stat().st_size})
        return files

    def delete_extra_knowledge_file(self, filename: str) -> bool:
        """Elimina un archivo de la carpeta global extra_knowledge/."""
        fpath = self._ensure_extra_knowledge_dir() / filename
        if fpath.exists():
            fpath.unlink()
            logger.info(f"Extra knowledge eliminado: {fpath}")
            return True
        return False

    # ──────────────────────────────────────────────────────────
    # Helpers internos
    # ──────────────────────────────────────────────────────────

    def _get_class_info(self, folder: Path, relative_path: str = None) -> Optional[dict]:
        """Construye el dict de información de una clase dado su folder y ruta relativa."""
        transcription_path = folder / "transcripcion.jsonl"
        summary_path = folder / "resumen.md"

        if not transcription_path.exists() and not summary_path.exists():
            return None

        created_at = datetime.fromtimestamp(folder.stat().st_mtime)

        segment_count = 0
        if transcription_path.exists():
            with open(transcription_path, "r", encoding="utf-8") as f:
                segment_count = sum(1 for line in f if line.strip())

        # Nombre para mostrar: parsear sufijo _YYYY-MM-DD_HH-MM si existe
        base_name = folder.name
        match = _DATE_SUFFIX_RE.match(base_name)
        if match:
            topic = match.group(1).replace("_", " ")
            date_part = match.group(2)   # YYYY-MM-DD
            time_part = match.group(3).replace("-", ":")  # HH:MM
            try:
                d = datetime.strptime(date_part, "%Y-%m-%d")
                display_name = f"{topic} · {d.strftime('%d/%m/%Y')} {time_part}"
            except ValueError:
                display_name = base_name.replace("_", " ")
        else:
            display_name = base_name.replace("_", " ")

        class_id = relative_path if relative_path else base_name
        folder_path = "/".join(class_id.split("/")[:-1]) if "/" in class_id else ""

        # Nombre personalizado si el usuario lo cambió
        custom_name_path = folder / "nombre.txt"
        if custom_name_path.exists():
            custom = custom_name_path.read_text(encoding="utf-8").strip()
            if custom:
                if match:
                    try:
                        d = datetime.strptime(match.group(2), "%Y-%m-%d")
                        time_part = match.group(3).replace("-", ":")
                        display_name = f"{custom} · {d.strftime('%d/%m/%Y')} {time_part}"
                    except ValueError:
                        display_name = custom
                else:
                    display_name = custom

        return {
            "id": class_id,
            "name": display_name,
            "folder_path": folder_path,
            "created_at": created_at.isoformat(),
            "created_at_formatted": created_at.strftime("%d/%m/%Y %H:%M"),
            "has_transcription": transcription_path.exists(),
            "has_summary": summary_path.exists(),
            "has_slides": (folder / "slides.md").exists(),
            "segment_count": segment_count,
            "path": str(folder)
        }

    @staticmethod
    def _strip_accents(text: str) -> str:
        """Elimina tildes/diacríticos de un texto, preservando caracteres base."""
        nfkd = unicodedata.normalize("NFKD", text)
        return "".join(c for c in nfkd if not unicodedata.combining(c))

    @staticmethod
    def _sanitize_folder_name(name: str) -> str:
        """Sanitiza un nombre de carpeta: sin tildes, sin espacios, sin caracteres especiales."""
        # Eliminar tildes/diacríticos para evitar problemas en rutas del sistema
        nfkd = unicodedata.normalize("NFKD", name)
        name = "".join(c for c in nfkd if not unicodedata.combining(c))
        name = name.replace(" ", "_").replace("-", "_")
        sanitized = "".join(c for c in name if c.isalnum() or c == "_")
        if len(sanitized) > 60:
            sanitized = sanitized[:60]
        return sanitized or "Clase_Sin_Nombre"

    @staticmethod
    def _sanitize_filename(filename: str) -> str:
        """Sanitiza un nombre de archivo: sin tildes en ruta, añade timestamp."""
        path = Path(filename)
        name = path.stem
        ext = path.suffix
        # Eliminar tildes/diacríticos
        nfkd = unicodedata.normalize("NFKD", name)
        name = "".join(c for c in nfkd if not unicodedata.combining(c))
        name = "".join(c for c in name if c.isalnum() or c in "._- ")
        name = name.replace(" ", "_")
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{name}_{timestamp}{ext}"
